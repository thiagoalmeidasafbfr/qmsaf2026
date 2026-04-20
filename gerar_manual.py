#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Gera MANUAL_GESTOR.docx e MANUAL_GESTOR.pdf a partir do conteúdo do manual."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy, os

OUT_DOCX = "MANUAL_GESTOR.docx"

# ── Cores ─────────────────────────────────────────────────────────────────────
C_BLACK   = RGBColor(0x0F, 0x17, 0x2A)
C_TITLE   = RGBColor(0x00, 0x22, 0x4E)   # azul escuro SAF
C_ACCENT  = RGBColor(0x00, 0x5B, 0xB5)   # azul
C_GRAY    = RGBColor(0x47, 0x55, 0x69)
C_LIGHT   = RGBColor(0xF1, 0xF5, 0xF9)
C_WHITE   = RGBColor(0xFF, 0xFF, 0xFF)
C_GREEN   = RGBColor(0x16, 0x6A, 0x34)
C_NOTE_BG = RGBColor(0xEF, 0xF6, 0xFF)

FONT = "Calibri"


# ── Helpers ───────────────────────────────────────────────────────────────────
def set_cell_bg(cell, rgb: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    hex_color = f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), kwargs.get("val", "single"))
        border.set(qn("w:sz"), kwargs.get("sz", "4"))
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), kwargs.get("color", "D1D5DB"))
        tcBorders.append(border)
    tcPr.append(tcBorders)


def para_space(doc, before=0, after=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = FONT
    run.font.bold = True
    if level == 1:
        run.font.size = Pt(15)
        run.font.color.rgb = C_TITLE
    else:
        run.font.size = Pt(12)
        run.font.color.rgb = C_ACCENT
    return p


def add_body(doc, text, bold_parts=None):
    """Add paragraph, optionally bold-ifying tokens wrapped in **."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    _add_rich_run(p, text)
    return p


def _add_rich_run(p, text):
    """Parse **bold** and `code` markers and add styled runs."""
    import re
    tokens = re.split(r'(\*\*[^*]+\*\*|`[^`]+`)', text)
    for tok in tokens:
        if tok.startswith("**") and tok.endswith("**"):
            r = p.add_run(tok[2:-2])
            r.bold = True
            r.font.name = FONT
            r.font.size = Pt(10.5)
        elif tok.startswith("`") and tok.endswith("`"):
            r = p.add_run(tok[1:-1])
            r.font.name = "Courier New"
            r.font.size = Pt(9.5)
            r.font.color.rgb = RGBColor(0xDC, 0x26, 0x26)
        else:
            r = p.add_run(tok)
            r.font.name = FONT
            r.font.size = Pt(10.5)
            r.font.color.rgb = C_BLACK


def add_numbered_item(doc, number, text, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.2 if indent else 0.6)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(f"{number}.  ")
    r.font.name = FONT
    r.font.size = Pt(10.5)
    r.font.bold = True
    r.font.color.rgb = C_ACCENT
    _add_rich_run(p, text)
    return p


def add_bullet(doc, text, indent_cm=0.6):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(indent_cm + 0.4)
    p.paragraph_format.first_line_indent = Cm(-0.4)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("•  ")
    r.font.name = FONT
    r.font.color.rgb = C_ACCENT
    r.font.size = Pt(10.5)
    _add_rich_run(p, text)
    return p


def add_note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single"); left.set(qn("w:sz"), "6")
    left.set(qn("w:space"), "10"); left.set(qn("w:color"), "005BB5")
    pBdr.append(left)
    pPr.append(pBdr)
    _add_rich_run(p, text)
    # italic everything
    for run in p.runs:
        run.italic = True
    return p


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # header row
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        set_cell_bg(cell, C_TITLE)
        cell.paragraphs[0].clear()
        run = cell.paragraphs[0].add_run(h)
        run.font.name = FONT
        run.font.size = Pt(9.5)
        run.font.bold = True
        run.font.color.rgb = C_WHITE
        cell.paragraphs[0].paragraph_format.space_before = Pt(4)
        cell.paragraphs[0].paragraph_format.space_after = Pt(4)

    # data rows
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        bg = RGBColor(0xF8, 0xFA, 0xFC) if ri % 2 == 0 else C_WHITE
        for ci, cell_text in enumerate(row_data):
            cell = row.cells[ci]
            set_cell_bg(cell, bg)
            cell.paragraphs[0].clear()
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            _add_rich_run(p, str(cell_text))
            for run in p.runs:
                run.font.size = Pt(9.5)

    # column widths
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


def add_divider(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1"); bottom.set(qn("w:color"), "CBD5E1")
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)


# ── Document ──────────────────────────────────────────────────────────────────
doc = Document()

# page margins
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)
section.top_margin    = Cm(2.2)
section.bottom_margin = Cm(2.2)

# default paragraph style
style = doc.styles["Normal"]
style.font.name = FONT
style.font.size = Pt(10.5)
style.font.color.rgb = C_BLACK

# ── CAPA ──────────────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(60)
p.paragraph_format.space_after  = Pt(6)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("SAF BOTAFOGO")
r.font.name = FONT; r.font.size = Pt(11); r.font.bold = True
r.font.color.rgb = C_ACCENT; r.font.all_caps = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(4)
r = p.add_run("Guia do Gestor")
r.font.name = FONT; r.font.size = Pt(28); r.font.bold = True
r.font.color.rgb = C_TITLE

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(2)
r = p.add_run("Portal QM — Quadro Móvel")
r.font.name = FONT; r.font.size = Pt(13)
r.font.color.rgb = C_GRAY

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(40)
r = p.add_run("Para: Gestores de setor  ·  Acesso restrito ao seu setor")
r.font.name = FONT; r.font.size = Pt(10); r.italic = True
r.font.color.rgb = C_GRAY

doc.add_page_break()

# ── SUMÁRIO MANUAL ─────────────────────────────────────────────────────────────
add_heading(doc, "Sumário")
sections_toc = [
    ("1.", "Fazendo login"),
    ("2.", "Sua tela inicial"),
    ("3.", "Adicionar um registro manualmente"),
    ("4.", "Importar registros em lote (planilha)"),
    ("5.", "Ver e filtrar o histórico"),
    ("6.", "Exportar o histórico"),
    ("7.", "Excluir registros"),
    ("8.", "Redefinir sua senha"),
    ("9.", "Entendendo a classificação QM"),
]
for num, title in sections_toc:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.left_indent  = Cm(0.4)
    r = p.add_run(f"{num}  ")
    r.font.name = FONT; r.font.bold = True; r.font.color.rgb = C_ACCENT; r.font.size = Pt(10.5)
    r2 = p.add_run(title)
    r2.font.name = FONT; r2.font.size = Pt(10.5); r2.font.color.rgb = C_BLACK

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 1. FAZENDO LOGIN
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "1. Fazendo login")
add_numbered_item(doc, 1, "Abra o portal no navegador.")
add_numbered_item(doc, 2, "Digite seu **e-mail** e **senha**.")
add_numbered_item(doc, 3, "Clique em **Entrar**.")
add_body(doc, "Você será direcionado diretamente para o seu painel.")
add_note(doc, "Caso o login não funcione, clique em **\"Esqueceu a senha?\"** e siga as instruções enviadas para o seu e-mail.")
add_divider(doc)

# ══════════════════════════════════════════════════════════════════════════════
# 2. TELA INICIAL
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "2. Sua tela inicial")
add_body(doc, "Após o login, você verá o **Painel do Gestor** com os seguintes elementos:")
add_table(doc,
    ["Elemento", "O que mostra"],
    [
        ["**Total de Registros**", "Todos os registros que você já lançou"],
        ["**Registros do Mês**",   "Apenas os registros do mês atual"],
        ["**Novo Registro**",      "Abre o formulário para lançar um registro manualmente"],
        ["**Importar Planilha**",  "Permite importar vários registros de uma vez via Excel"],
        ["**Ver Histórico**",      "Abre a lista de todos os seus registros"],
    ],
    col_widths=[5.5, 10.5]
)
add_divider(doc)

# ══════════════════════════════════════════════════════════════════════════════
# 3. ADICIONAR REGISTRO MANUALMENTE
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "3. Adicionar um registro manualmente")
add_numbered_item(doc, 1, "Clique no botão **Novo Registro**.")
add_numbered_item(doc, 2, "Preencha o formulário conforme descrito abaixo.")

add_heading(doc, "Campos do formulário", level=2)

fields = [
    ("Tipo de Evento",
     "Escolha entre `Jogo` (para jogos das categorias do clube) ou `Outros` (demais eventos)."),
    ("Nome do Evento",
     "Escreva o nome do evento. Ex.: `Botafogo x Flamengo`, `Treino Sub-20`."),
    ("Data do Evento",
     "Selecione a data em que o evento acontece (não a data de hoje)."),
    ("Horário do Evento",
     "Escolha o horário de início. Disponível em intervalos de **15 minutos** (19:00, 19:15, 19:30...)."),
    ("Categoria do Jogo",
     "Somente para tipo **Jogo**. Opções: Futebol Masculino Profissional, Sub-20, Sub-17, Sub-15, Demais Categorias Masculino, Futebol Profissional Feminino, Futebol de Base Feminino."),
    ("Tipo QM",
     "Selecione `A` ou `B`. O sistema calcula a classificação final (A, B, C ou D) automaticamente. Veja a seção 9."),
    ("Funcionário",
     "Selecione o nome na lista. Se não constar, selecione `Outros`."),
    ("Observações",
     "**Obrigatório** quando Tipo = `Outros` ou Funcionário = `Outros`. Nos demais casos é opcional."),
]
for name, desc in fields:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(f"{name}:  ")
    r.font.name = FONT; r.font.bold = True; r.font.size = Pt(10.5); r.font.color.rgb = C_TITLE
    _add_rich_run(p, desc)

add_numbered_item(doc, 3, "Clique em **Salvar Registro**.")
add_numbered_item(doc, 4, "Uma mensagem de confirmação aparecerá e o registro será listado no histórico.")
add_divider(doc)

# ══════════════════════════════════════════════════════════════════════════════
# 4. IMPORTAR VIA PLANILHA
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "4. Importar registros em lote (planilha)")
add_body(doc, "Use essa opção quando precisar lançar vários registros de uma vez.")

add_heading(doc, "Passo 1 — Baixar o modelo", level=2)
add_numbered_item(doc, 1, "Clique em **Importar Planilha**.")
add_numbered_item(doc, 2, "Clique em **Baixar Modelo** para salvar o arquivo `modelo_quadro_movel.xlsx`.")
add_numbered_item(doc, 3, "Abra no **Excel** ou **Google Planilhas**.")
add_body(doc, "O arquivo contém três abas:")
add_bullet(doc, "**Registros** — onde você preenche os dados")
add_bullet(doc, "**Listas** — referências (funcionários, horários, categorias)")
add_bullet(doc, "**Tutorial** — explicação de cada coluna")

add_heading(doc, "Passo 2 — Preencher a planilha", level=2)
add_body(doc, "Preencha a aba **Registros** a partir da **linha 2**. Não altere ou apague a linha 1 (cabeçalhos).")
add_table(doc,
    ["Coluna", "Obrigatório?", "O que preencher"],
    [
        ["Tipo de Evento",    "Sim",         "`Jogo` ou `Outros`"],
        ["Nome do Evento",    "Sim",         "Texto livre"],
        ["Data do Evento",    "Sim",         "`DD/MM/AAAA` — ex.: `16/04/2026`"],
        ["Horário do Evento", "Sim",         "`HH:MM` — ex.: `19:00` (múltiplos de 15 min)"],
        ["Categoria do Jogo", "Só p/ Jogo",  "Uma das categorias da aba Listas"],
        ["Tipo QM",           "Sim",         "Apenas `A` ou `B`"],
        ["Funcionário",       "Sim",         "Nome exato ou `Outros`"],
        ["Observações",       "Condicional", "Obrigatório se Tipo = Outros ou Funcionário = Outros"],
    ],
    col_widths=[4.5, 3.2, 8.3]
)
add_note(doc, "Para eventos do tipo `Outros`, a coluna **Categoria do Jogo deve ficar em branco**. Os nomes dos funcionários devem ser **exatamente iguais** ao cadastro (use a aba Listas para consultar).")

add_heading(doc, "Passo 3 — Importar", level=2)
for i, step in enumerate([
    "Salve o arquivo como `.xlsx`.",
    "No portal, clique em **Importar Planilha**.",
    "Arraste o arquivo para a área indicada ou clique para selecionar.",
    "O sistema verificará cada linha. Erros serão exibidos com o número da linha e o motivo.",
    "Corrija os erros na planilha e importe novamente se necessário.",
    "Se houver duplicatas (mesma data + mesmo funcionário), o sistema avisará — você escolhe se ignora ou importa assim mesmo.",
], start=1):
    add_numbered_item(doc, i, step)

add_heading(doc, "Erros comuns", level=2)
add_table(doc,
    ["Mensagem de erro", "O que causou", "Como corrigir"],
    [
        ["`Tipo QM inválido`",         "Valor diferente de A ou B",          "Preencha apenas com `A` ou `B`"],
        ["`Formato de data inválido`", "Data fora do padrão",                "Use `DD/MM/AAAA` — ex.: `16/04/2026`"],
        ["`Categoria obrigatória`",    "Tipo = Jogo sem categoria",           "Preencha a Categoria do Jogo"],
        ["`Observações obrigatórias`", "Tipo = Outros sem observação",        "Preencha as Observações"],
        ["`Funcionário não encontrado`","Nome diferente do cadastro",         "Consulte a aba Listas para o nome exato"],
        ["`Cabeçalho ausente`",        "Coluna renomeada ou apagada",         "Use o modelo original sem alterar a linha 1"],
    ],
    col_widths=[5, 5, 6]
)
add_divider(doc)

# ══════════════════════════════════════════════════════════════════════════════
# 5. HISTÓRICO
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "5. Ver e filtrar o histórico")
add_numbered_item(doc, 1, "Clique em **Ver Histórico** no painel.")
add_numbered_item(doc, 2, "Você verá todos os registros do(s) seu(s) setor(es).")

add_heading(doc, "Filtros disponíveis", level=2)
add_body(doc, "Use os campos de busca no topo da tabela para filtrar por:")
for f in ["Nome do evento", "Nome do funcionário", "Categoria do jogo", "Período (data inicial e data final)", "Observações"]:
    add_bullet(doc, f)
add_body(doc, "Os filtros são aplicados automaticamente conforme você digita.")

add_heading(doc, "Colunas exibidas na tabela", level=2)
add_table(doc,
    ["Coluna", "Descrição"],
    [
        ["Data Criação", "Quando o registro foi lançado no sistema"],
        ["Evento",       "Nome do evento"],
        ["Data Evento",  "Data em que o evento ocorre"],
        ["Horário",      "Horário do evento"],
        ["Categoria",    "Categoria do jogo"],
        ["QM",           "Classificação calculada (A, B, C ou D)"],
        ["Funcionário",  "Nome do funcionário"],
        ["Setor",        "Setor do funcionário"],
        ["Observações",  "Observações inseridas"],
    ],
    col_widths=[4, 12]
)
add_divider(doc)

# ══════════════════════════════════════════════════════════════════════════════
# 6. EXPORTAR
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "6. Exportar o histórico")
add_numbered_item(doc, 1, "Acesse **Ver Histórico**.")
add_numbered_item(doc, 2, "Aplique filtros se quiser exportar apenas um período ou funcionário específico.")
add_numbered_item(doc, 3, "Clique em **Exportar CSV** ou **Exportar XLSX**.")
add_numbered_item(doc, 4, "O arquivo `historico_registros.xlsx` (ou `.csv`) será baixado automaticamente.")
add_note(doc, "A exportação considera **apenas os registros visíveis** após os filtros. Para exportar tudo, limpe os filtros antes de exportar.")
add_divider(doc)

# ══════════════════════════════════════════════════════════════════════════════
# 7. EXCLUIR
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "7. Excluir registros")

add_heading(doc, "Excluir um registro", level=2)
add_body(doc, "Clique no ícone de lixeira ao lado do registro e confirme.")

add_heading(doc, "Excluir vários registros de uma vez", level=2)
add_numbered_item(doc, 1, "Marque as caixas de seleção dos registros que deseja excluir.")
add_numbered_item(doc, 2, "Clique em **Excluir Selecionados**.")
add_numbered_item(doc, 3, "Uma janela de confirmação mostrará os registros que serão removidos.")
add_numbered_item(doc, 4, "Clique em **Confirmar** para concluir.")
add_note(doc, "**Atenção:** a exclusão é permanente e não pode ser desfeita.")
add_divider(doc)

# ══════════════════════════════════════════════════════════════════════════════
# 8. REDEFINIR SENHA
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "8. Redefinir sua senha")
add_numbered_item(doc, 1, "Na tela de login, clique em **\"Esqueceu a senha?\"**.")
add_numbered_item(doc, 2, "Informe seu e-mail cadastrado.")
add_numbered_item(doc, 3, "Clique em **Enviar link de redefinição**.")
add_numbered_item(doc, 4, "Acesse o e-mail recebido e clique no link.")
add_numbered_item(doc, 5, "Cadastre sua nova senha.")
add_divider(doc)

# ══════════════════════════════════════════════════════════════════════════════
# 9. CLASSIFICAÇÃO QM
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "9. Entendendo a classificação QM")
add_body(doc, "Ao salvar um registro, o sistema calcula automaticamente a **classificação QM** (A, B, C ou D) e o **valor** com base em três fatores:")
add_numbered_item(doc, 1, "O **Tipo QM** informado (A ou B)")
add_numbered_item(doc, 2, "O **dia da semana** do evento")
add_numbered_item(doc, 3, "O **horário** do evento")

add_heading(doc, "Tabela de classificação", level=2)
add_table(doc,
    ["Tipo QM", "Condição", "Classificação", "Valor"],
    [
        ["A", "Fim de semana (sáb/dom) ou horário ≥ 21:00", "**A**", "R$ 400"],
        ["B", "Fim de semana (sáb/dom) ou horário ≥ 21:00", "**B**", "R$ 200"],
        ["A", "Dia útil (seg–sex) e horário < 21:00",       "**C**", "R$ 200"],
        ["B", "Dia útil (seg–sex) e horário < 21:00",       "**D**", "R$ 100"],
    ],
    col_widths=[2.5, 8, 3.5, 2]
)

add_heading(doc, "Exemplos práticos", level=2)
add_table(doc,
    ["Situação", "Tipo QM", "Resultado"],
    [
        ["Sexta-feira, 19:00",          "A", "C — R$ 200"],
        ["Sexta-feira, 21:00",          "A", "A — R$ 400"],
        ["Sábado, 16:00",               "B", "B — R$ 200"],
        ["Quarta-feira, 20:59",         "B", "D — R$ 100"],
        ["Domingo, qualquer horário",   "A", "A — R$ 400"],
    ],
    col_widths=[7, 3, 6]
)
add_note(doc, "O horário **21:00 em ponto** já conta como noturno (≥ 21:00 = classificação mais alta).")

add_divider(doc)

# rodapé final
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(16)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Em caso de dúvidas ou problemas de acesso, contate o administrador do sistema.")
r.font.name = FONT; r.font.size = Pt(9.5); r.italic = True; r.font.color.rgb = C_GRAY

# ── Salvar DOCX ───────────────────────────────────────────────────────────────
doc.save(OUT_DOCX)
print(f"✓ {OUT_DOCX} gerado com sucesso.")
